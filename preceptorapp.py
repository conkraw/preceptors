import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import io
import re
import json
import pytz
import firebase_admin
from firebase_admin import credentials, firestore
import openai
import zipfile
import docx
import random 
import numpy as np
from functools import reduce
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from contextlib import suppress

###ERROR LIKELY WITH FIRBEASE LINE 473... 

st.set_page_config(layout="wide")

def safe_check_and_add_record(record_id):
    try:
        return check_and_add_record(record_id)
    except Exception as e:
        st.error(f"Error processing record {record_id}: {e}")
        return True  # You might choose to exclude this record or handle it differently


# Define a function to group every two elements into a single name
def group_names(name_str):
    if isinstance(name_str, str):
        parts = [part.strip() for part in name_str.split(',')]
        return [f"{parts[i]}, {parts[i+1]}" for i in range(0, len(parts)-1, 2)]
    else:
        return []


def shade_cell(cell, shade_color="D3D3D3"):
    """
    Applies background shading to a table cell.
    shade_color should be a hex color code (without the '#').
    For example, 'D3D3D3' is a light gray.
    """
    # Get the cell properties
    tcPr = cell._tc.get_or_add_tcPr()
    # Look for an existing w:shd element
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        # If not found, create one
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    # Set the shading style
    # 'val' can be 'clear', 'solid', etc.
    # 'color' is usually 'auto' if you're only changing fill.
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shade_color)
    
def set_cell_border(cell, **kwargs):
    """
    Set cell borders. Accepts arguments:
    top, start, bottom, end, left, right in the form of:
    {"sz": "12", "val": "single", "color": "000000"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end', 'left', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcPr.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcPr.append(element)
                for key, value in edge_data.items():
                    element.set(qn('w:{}'.format(key)), value)

# A small helper to create a two-row table (header row + content row)



def set_table_width(table, width_in_inches):
    """
    Force the overall table width (in inches).
    """
    # Convert inches to twips (1 inch = 1440 twips)
    width_in_twips = int(width_in_inches * 1440)

    tblPr = table._tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(width_in_twips))
    tblW.set(qn('w:type'), 'dxa')  # 'dxa' means absolute width

def create_comment_table(document, header_text, content_text, table_width=6.14):
    table = document.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    
    # Disable autofit to help ensure the width is fixed
    table.autofit = False

    # Force the table to the desired width
    set_table_width(table, table_width)
    
    # 1) Header row
    hdr_cell = table.cell(0, 0)
    hdr_cell.text = header_text

    shade_cell(hdr_cell, "D3D3D3")  # Light gray
    
    for paragraph in hdr_cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

    # 2) Content row
    table.cell(1, 0).text = str(content_text)

    
def set_cell_width(cell, width_in_inches):
    # 1 inch = 1440 dxa (twips)
    width_dxa = int(width_in_inches * 1440)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """
    Set cell borders. Each argument (top/bottom/left/right) is a dict of border attributes:
      {
        "sz": "12",       # border width
        "val": "single",  # single, double, etc.
        "color": "000000" # hex color code (omit #)
      }
    Example usage:
      set_cell_border(cell,
                      top={"sz": "12", "val": "single", "color": "000000"},
                      left={"sz": "12", "val": "single", "color": "000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Map python-docx edge names to the w: tag names
    # "left" -> "w:left", etc.
    borders = {
        'top': 'w:top',
        'bottom': 'w:bottom',
        'left': 'w:left',
        'right': 'w:right'
    }
    
    for edge, edge_data in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if edge_data:
            tag = borders[edge]
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))



def generate_spotlight_summary(strengths_preceptor, Evaluator):
    prompt = f"""
    You are an expert in pediatric medical education.

    Based on the following strengths feedback for {Evaluator}:
    {strengths_preceptor}

    Provide a concise summary of why this preceptor deserves to be in the spotlight.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in pediatric medical education."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=150,
    )
    return response['choices'][0]['message']['content'].strip()


def check_and_add_record(record_id):
    # Ensure the record_id is a string
    record_id_str = str(record_id)
    doc_ref = db.collection("preceptors").document(record_id_str)
    
    # If the record does not exist, add it
    if not doc_ref.get().exists:
        # Optionally, you can add additional data (like a timestamp) in the document
        doc_ref.set({"processed": True})
        return False  # Indicates the record was not previously processed
    else:
        return True  # Indicates the record already exists

        
def strengths(strengths_preceptor, Evaluator):
    prompt = f"""
    You are an expert in pediatric medical education.

    {Evaluator} received the following feedback regarding their performance as a preceptor in a pediatric clerkship:
    {strengths_preceptor}

    Please provide a concise summary of {Evaluator}'s strengths.
    In your summary, refer to the individual by name (using their first and last name and/or “Dr. Lastname”) when describing actions or behaviors.
    Assume that the feedback pertains exclusively to one individual.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in pediatric medical education."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500,
    )
    return response['choices'][0]['message']['content'].strip()

def improvement(improvement_preceptor, Evaluator):
    prompt = f"""
    You are an expert in pediatric medical education.

    {Evaluator} received the following feedback regarding opportunities for improvement as a preceptor in a pediatric clerkship:
    {improvement_preceptor}

    Please provide a concise summary of {Evaluator}'s opportunities for improvement.
    In your summary, refer to the individual by name (using their first and last name and/or “Dr. Lastname”) when describing actions or behaviors.
    Assume that the feedback pertains exclusively to one individual.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in pediatric medical education."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500,
    )
    return response['choices'][0]['message']['content'].strip()

def summarize_feedback(comments):
    # Example prompt that avoids identifying individual students
    prompt = f"""
    You are an expert in pediatric medical education.

    The following comments summarize documentation issues observed among multiple students:

    \"\"\"{comments}\"\"\"

    Provide a concise, high-level summary of the primary documentation challenges 
    these students faced. Do not reference individual students. 
    Focus on how the preceptor can improve future students' documentation skills.
    """

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in pediatric medical education."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500
    )
    return response['choices'][0]['message']['content'].strip()
    
########################################
# 1) OPENAI & FIREBASE SETUP
########################################

openai.api_key = st.secrets["openai"]["api_key"]

firebase_creds = st.secrets["firebase_service_account"].to_dict()
if not firebase_admin._apps:
    cred = credentials.Certificate(firebase_creds)
    firebase_admin.initialize_app(cred)

db = firestore.client()

########################################
# 2) DISPLAY UPLOAD SECTIONS
########################################

st.title("Faculty Analysis Report & Evaluation Due Dates Upload")
st.subheader("Please Upload in order... ")

col1, col2, col3 = st.columns(3)

with col1:
    st.subheader("Redcap Metrics")
    # Embed the website so users can view it directly
    st.markdown("[REDCAP Link Raw File](https://redcap.ctsi.psu.edu/redcap_v14.5.43/DataExport/index.php?pid=16813&report_id=61309)")
    redcapmetrics = st.file_uploader("Upload REDCAP Report", type=["pdf", "docx", "csv"])

with col2:
    st.subheader("Student Clinical Assessment Forms")
    st.markdown("[Student Performance Report](https://oasis.pennstatehealth.net/admin/course/e_manage/student_performance/setup_analysis_report.html)")
    # File uploader for evaluation due dates (adjust allowed types as needed)
    evaluation_due_dates_file = st.file_uploader("Upload Evaluation Due Dates", type=["csv", "xlsx", "pdf"])

with col3:
    st.subheader("Student Evaluations of Preceptors")
    # Embed the website so users can view it directly
    st.markdown("[Faculty Evaluations](https://oasis.pennstatehealth.net/admin/course/e_manage/faculty/setup_analysis_report.html)")
    analysis_report_file = st.file_uploader("Upload Analysis Report", type=["pdf", "docx", "csv"])
    




    
if redcapmetrics is not None:
    try:
        # Determine the file type and load accordingly
        if redcapmetrics.name.endswith("csv"):
            dfe = pd.read_csv(redcapmetrics)
        elif redcapmetrics.name.endswith("xlsx"):
            dfe = pd.read_excel(redcapmetrics)
        # Display the DataFrame in the app
        
        dff = dfe 
        dfg = dfe 
        dfh = dfe 
        dfi = dfe
        dfj = dfe 
        
        dfe = dfe.dropna(subset=['oasis_cas'])
        dfe['corrected_preceptors'] = dfe['oasis_cas'].apply(group_names)
        # Explode to individual rows preserving full names correctly
        dfe = dfe.explode('corrected_preceptors')
        # Count occurrences correctly for each full preceptor name by unique record_id
        dfe = dfe.groupby('corrected_preceptors')['record_id'].nunique().reset_index()
        dfe.columns = ['corrected_preceptors', 'student_matches']
        # Sort results
        dfe = dfe.sort_values(by='student_matches', ascending=False)

        dff = dff.dropna(subset=['oasis_cas'])
        # Combine 'week' columns into one using melt
        df_melted = dff.melt(id_vars=['record_id'], value_vars=['week1', 'week2', 'week3', 'week4'], var_name='week', value_name='week_preceptor')
        # Drop NaNs resulting from melting
        df_melted = df_melted.dropna(subset=['week_preceptor'])
        # Apply group_names function to melted data safely
        df_melted['corrected_preceptors'] = df_melted['week_preceptor'].apply(group_names)
        # Explode to individual rows per preceptor
        df_exploded = df_melted.explode('corrected_preceptors')
        # Clean whitespace
        df_exploded['corrected_preceptors'] = df_exploded['corrected_preceptors'].str.strip()
        # Count each occurrence (preceptor matched per student per week separately)
        student_assignments = df_exploded.groupby('corrected_preceptors').size().reset_index(name='student_assignments').sort_values(by='student_assignments', ascending=False)
        dff = student_assignments

        df = dfg.dropna(subset=['oasis_cas'])
        # Calculate student-level average scores (ignoring NaNs)
        df['average_prac_score'] = df[['total_prac_scorep_v1', 'total_prac_scorep_v2']].mean(axis=1, skipna=True)
        # Apply grouping of names
        df['corrected_preceptors'] = df['oasis_cas'].apply(group_names)
        # Explode dataframe (one preceptor per row)
        df_exploded = df.explode('corrected_preceptors')
        # Trim whitespace
        df_exploded['corrected_preceptors'] = df_exploded['corrected_preceptors'].str.strip()
        # Calculate average score per preceptor
        preceptor_avg_scores = df_exploded.groupby('corrected_preceptors')['average_prac_score'] \
                                 .mean().reset_index(name='average_prac_score')
        preceptor_avg_sorted = preceptor_avg_scores.sort_values(by='average_prac_score', ascending=False)
        dfg = preceptor_avg_sorted

        df = dfh.dropna(subset=['oasis_cas'])
        # Calculate student-level average scores (ignoring NaNs)
        df['average_doc_score'] = df[['scorep_v1', 'scorep_v2']].mean(axis=1, skipna=True)
        # Apply grouping of names
        df['corrected_preceptors'] = df['oasis_cas'].apply(group_names)
        # Explode dataframe (one preceptor per row)
        df_exploded = df.explode('corrected_preceptors')
        # Trim whitespace
        df_exploded['corrected_preceptors'] = df_exploded['corrected_preceptors'].str.strip()
        # Calculate average score per preceptor
        preceptor_avg_scores = df_exploded.groupby('corrected_preceptors')['average_doc_score'] \
                                 .mean().reset_index(name='average_doc_score')
        preceptor_avg_sorted = preceptor_avg_scores.sort_values(by='average_doc_score', ascending=False)
        dfh = preceptor_avg_sorted

        df = dfi.dropna(subset=['oasis_cas'])
        # Calculate student-level average scores (ignoring NaNs)
        df['average_nbme'] = df[['nbme']].mean(axis=1, skipna=True)
        # Apply grouping of names
        df['corrected_preceptors'] = df['oasis_cas'].apply(group_names)
        # Explode dataframe (one preceptor per row)
        df_exploded = df.explode('corrected_preceptors')
        # Trim whitespace
        df_exploded['corrected_preceptors'] = df_exploded['corrected_preceptors'].str.strip()
        # Calculate average score per preceptor
        preceptor_avg_scores = df_exploded.groupby('corrected_preceptors')['average_nbme'] \
                                 .mean().reset_index(name='average_nbme')
        preceptor_avg_sorted = preceptor_avg_scores.sort_values(by='average_nbme', ascending=False)
        dfi = preceptor_avg_sorted

        df = dfj 
        df['combined_comments'] = (df[['doccomment_v1', 'doccomment_v2']].apply(lambda row: ' '.join(row.dropna().astype(str)).strip(), axis=1))
        df['combined_comments'] = df['combined_comments'].replace(r'^\s*$', np.nan, regex=True)
        df = df.dropna(subset=['combined_comments']).copy()
        df['corrected_preceptors'] = df['oasis_cas'].apply(group_names)
        df_exploded = df.explode('corrected_preceptors')
        df_grouped = df_exploded.groupby('corrected_preceptors')['combined_comments'].apply(lambda rows: ' '.join(rows)).reset_index(name='all_comments')

        #######################AI DOCUMENTATION SUMMARY#######################
        #df_grouped['documentation_summary'] = df_grouped['all_comments'].apply(summarize_feedback)

        df_grouped['documentation_summary'] = df_grouped['all_comments'].astype(str)
        ######################################################################
        
        final_df = df_grouped[['corrected_preceptors', 'documentation_summary']]

        dfj = final_df

        df_list = [dfe, dff, dfg, dfh, dfi, dfj]
        df_final = reduce(lambda left, right: pd.merge(left, right, on='corrected_preceptors', how='outer'), df_list)
        df_final.rename(columns={'corrected_preceptors': 'Evaluator'}, inplace=True)
        groupedx = df_final.reset_index()
        # Now, if you still need to set it as the index for further operations:
        groupedx = groupedx.set_index('Evaluator')
        st.write("RedCap Metrics Uploaded")

    except Exception as e:
        st.error(f"Error loading file: {e}")

if evaluation_due_dates_file is not None:
    try:
        # Determine the file type and load accordingly
        if evaluation_due_dates_file.name.endswith("csv"):
            dfe = pd.read_csv(evaluation_due_dates_file)
        elif evaluation_due_dates_file.name.endswith("xlsx"):
            dfe = pd.read_excel(evaluation_due_dates_file)
        # Display the DataFrame in the app
        dfe = dfe.loc[dfe['Location'] != "LIC - Kaiser Permanente"]
        dfe = dfe[['Evaluator', 'Submit Date', 'End Date']]

        # Convert the date columns to datetime format
        dfe['Submit Date'] = pd.to_datetime(dfe['Submit Date'])
        dfe['End Date'] = pd.to_datetime(dfe['End Date'])
        
        # Calculate the difference in days between Submit Date and End Date
        dfe['diff_days'] = (dfe['Submit Date'] - dfe['End Date']).dt.days
        
        # Create a boolean flag for evaluations that are less than or equal to 14 days
        dfe['on_time'] = dfe['diff_days'] <= 14
        
        grouped = dfe.groupby('Evaluator').agg(total_evaluations=('Evaluator', 'size'),on_time_evaluations=('on_time', 'sum'))
        
        # Calculate the percentage of on-time evaluations per evaluator
        grouped['percentage_on_time'] = ((grouped['on_time_evaluations'] / grouped['total_evaluations']) * 100).round(1)

        grouped = grouped.reset_index()
        # Now, if you still need to set it as the index for further operations:
        grouped = grouped.set_index('Evaluator')
        st.write("Evaluation Due Dates Uploaded")
        

    except Exception as e:
        st.error(f"Error loading file: {e}")

if analysis_report_file is not None:
    try:
        # Determine the file type and load accordingly
        if analysis_report_file.name.endswith("csv"):
            dfa = pd.read_csv(analysis_report_file)
        elif analysis_report_file.name.endswith("xlsx"):
            dfa = pd.read_excel(analysis_report_file)
        # Display the DataFrame in the app

        ###########################################################################################
        # First, filter the DataFrame based on Firebase:
        #dfa = dfa[~dfa["Form Record"].apply(safe_check_and_add_record)]
        ###########################################################################################
        
        selected_indices = [4, 5, 16, 19, 23, 27, 30, 34, 37, 41, 44, 48, 51, 55, 58, 62,65, 69, 72, 76, 79, 83, 86, 90, 93, 97, 100, 104, 107, 111, 114, 118, 121, 125, 128, 132, 135, 139, 143, 146, 147, 153, 154]
        dfa = dfa.iloc[:, selected_indices]

        df = dfa.copy()
        rename_mapping = {}
        # Loop through columns to find ones with the pattern "<number> Question"
        for col in df.columns:
            m = re.match(r'^(\d+)\s+Question$', col)
            if m:
                num = m.group(1)
                # Try different possible suffixes for the corresponding answer column
                for suffix in ["Multiple Choice Value", "Multiple Choice Label", "Answer text"]:
                    answer_col = f"{num} {suffix}"
                    if answer_col in df.columns:
                        # Use the question text from the first row of the question column
                        question_text = df[col].iloc[0] if not df[col].empty else col
                        rename_mapping[answer_col] = question_text
                        break  # stop after the first matching answer column

        # Rename the answer columns with the question text
        df.rename(columns=rename_mapping, inplace=True)

        # Optionally, remove the question columns if they are no longer needed
        question_columns = [col for col in df.columns if re.match(r'^\d+\s+Question$', col)]
        df.drop(columns=question_columns, inplace=True)

        df['Rotation Period'] = pd.to_datetime(df.iloc[:, 0], errors='coerce').dt.strftime('%B %Y')
    
        # --- Step 2. Drop Unneeded Date Columns ---
        # Drop 'Start Date' (col 0) and 'End Date' (col 1)
        df.drop(df.columns[[0, 1]], axis=1, inplace=True)
        
        # --- Step 3. Drop the Unwanted Time Column ---
        df.drop(df.columns[19], axis=1, inplace=True)
        
        # --- Step 4. Rename the Text Columns ---
        # After dropping one column, the strengths and improvement columns are now at positions 19 and 20 respectively.
        df.columns.values[19] = "strengths_preceptor"
        df.columns.values[20] = "improvement_preceptor"
        
        # --- Step 5. Convert Evaluation Score Columns to Numeric ---
        # The evaluation questions are now columns 3 to 18. We refer to them by position.
        score_cols = df.columns[3:19]
        df[score_cols] = df[score_cols].apply(pd.to_numeric, errors='coerce')
        
        # --- Step 6. Group by Evaluator (and Rotation Period) and Aggregate ---
        # If an evaluator appears more than once for a given rotation period, we’ll compute the mean for each question.
        # We also include Evaluator Email and Form Record as identifying columns.
        group_cols = ["Evaluator", "Evaluator Email", "Rotation Period", "Form Record"]
        
        # For the score columns, take the mean.
        agg_dict = {col: 'mean' for col in score_cols}
        # For text responses, combine unique responses (ignoring NaN) using a separator.
        agg_dict["strengths_preceptor"] = lambda x: ' | '.join(x.dropna().unique())
        agg_dict["improvement_preceptor"] = lambda x: ' | '.join(x.dropna().unique())
        
        # Group the DataFrame by the identifying columns and aggregate
        df_grouped = df.groupby(group_cols, as_index=False).agg(agg_dict)
        
        # --- (Optional) Reorder Columns for Clarity ---
        ordered_columns = group_cols + list(score_cols) + ["strengths_preceptor", "improvement_preceptor"]
        df_grouped = df_grouped[ordered_columns]
        
        # Display the final aggregated DataFrame in your Streamlit app
        #st.dataframe(df_grouped)

        final_group_cols = ["Evaluator", "Evaluator Email"]
        
        agg_funcs = {
            # For Rotation Period: convert each unique item to a string before joining with a comma
            "Rotation Period": lambda x: ", ".join([str(item) for item in sorted(set(x.dropna()))]),
            # For strengths and improvements: convert each unique item to a string before joining with a newline
            "strengths_preceptor": lambda x: "\n".join([str(item) for item in x.dropna().unique()]),
            "improvement_preceptor": lambda x: "\n".join([str(item) for item in x.dropna().unique()])
        }
        
        # Identify the evaluation score columns (all numeric columns not already in our aggregation)
        score_columns = [col for col in df_grouped.columns 
                         if col not in ["Evaluator", "Evaluator Email", "Rotation Period", 
                                        "strengths_preceptor", "improvement_preceptor", "Form Record"]]
        
        # For each score column, take the mean (this will average the score for each question)
        for col in score_columns:
            agg_funcs[col] = "mean"
        
        # Optionally, aggregate the Form Record if needed (or drop it)
        if "Form Record" in df_grouped.columns:
            agg_funcs["Form Record"] = lambda x: ", ".join([str(item) for item in x.dropna().unique()])
            
        df_grouped["num_evaluations"] = 1

        # Extend your aggregation dictionary to sum the number of evaluations
        agg_funcs["num_evaluations"] = "sum"
        
        # Group by Evaluator and Evaluator Email using the updated aggregation functions
        final_group_cols = ["Evaluator", "Evaluator Email"]
        df_final = df_grouped.groupby(final_group_cols, as_index=False).agg(agg_funcs)
        st.write("Preceptor Evaluation Pre AI")
        #############################################################################################################################
        df_final["strengths_summary"] = "test"
        df_final["improvement_summary"] = "test"
        
        #df_final["strengths_summary"] = df_final.apply(lambda row: strengths(row["strengths_preceptor"], row["Evaluator"]), axis=1)
        #df_final["improvement_summary"] = df_final.apply(lambda row: improvement(row["improvement_preceptor"], row["Evaluator"]), axis=1)
        #############################################################################################################################
        
        # Map the values to df_final
        df_final['total_evaluations'] = df_final['Evaluator'].map(grouped['total_evaluations'])
        df_final['percentage_on_time'] = df_final['Evaluator'].map(grouped['percentage_on_time'])
        
        df_final['student_matches'] = df_final['Evaluator'].map(groupedx['student_matches'])
        df_final['student_assignments'] = df_final['Evaluator'].map(groupedx['student_assignments'])
        df_final['average_prac_score'] = df_final['Evaluator'].map(groupedx['average_prac_score'])
        df_final['average_doc_score'] = df_final['Evaluator'].map(groupedx['average_doc_score'])
        df_final['average_nbme'] = df_final['Evaluator'].map(groupedx['average_nbme'])
        df_final['documentation_summary'] = df_final['Evaluator'].map(groupedx['documentation_summary'])
        
        # Display the final aggregated DataFrame with the count of evaluations
        st.dataframe(df_final)
        
        with suppress(NameError):
            if "df_final_summarized" not in st.session_state:
                with st.spinner("Summarizing strengths…"):
                    df_final["strengths_summary"] = df_final.apply(
                        lambda row: strengths(row["strengths_preceptor"], row["Evaluator"]), axis=1
                    )
                with st.spinner("Summarizing improvement…"):
                    df_final["improvement_summary"] = df_final.apply(
                        lambda row: improvement(row["improvement_preceptor"], row["Evaluator"]), axis=1
                    )
                with st.spinner("Summarizing documentation…"):
                    df_final["documentation_summary"] = df_final["documentation_summary"].apply(summarize_feedback)
                st.session_state["df_final_summarized"] = df_final.copy()
        
        # pull the cached one
        df_final = st.session_state["df_final_summarized"]
                
        # --- STEP 1: Identify Eligible Preceptors ---
        # Define the known text fields to identify numeric score columns.
        known_cols = {"Evaluator", "Evaluator Email", "Rotation Period", "strengths_preceptor", "improvement_preceptor", "strengths_summary", "improvement_summary", "num_evaluations", "Form Record", "total_evaluations", "percentage_on_time", "student_matches", "student_assignments", "average_prac_score", "average_doc_score", "average_nbme"}

        exclude_cols = {"total_evaluations", "percentage_on_time", "num_evaluations", "student_matches", "student_assignments", "average_prac_score", "average_doc_score", "average_nbme"}
        score_cols = [col for col in df_final.columns if pd.api.types.is_numeric_dtype(df_final[col]) and col not in exclude_cols]
        
        # Filter for eligible preceptors: every evaluation score must be 4.5 or above.
        eligible_df = df_final[df_final[score_cols].ge(4.5).all(axis=1)].copy()

        ###########################################################################################
        # Retrieve already spotlighted evaluators from Firebase.
        spotlight_docs = db.collection("spotlight").stream()
        spotlight_evaluators = {doc.to_dict().get("Evaluator") for doc in spotlight_docs}        
        # Exclude evaluators already in the spotlight.
        eligible_df = eligible_df[~eligible_df["Evaluator"].isin(spotlight_evaluators)]
        ###########################################################################################
        
        if eligible_df.empty:
            st.info("No eligible preceptors found for the spotlight this month.")
        else:
            ########################################################
            # --- STEP 2: Randomly Select a Spotlight Candidate ---#
            #selected_candidate = eligible_df.sample(n=1).iloc[0]  #
            ########################################################
            
            df_final = st.session_state["df_final_summarized"]
            
            # 0) Initialize our session‑state defaults exactly once
            if "_prev_preceptor" not in st.session_state:
                st.session_state["_prev_preceptor"] = None
            if "spotlight_reason" not in st.session_state:
                st.session_state["spotlight_reason"] = ""   # empty string, not undefined
            
            # 1) Let user pick someone
            preceptor_list = eligible_df["Evaluator"].unique()
            selected_preceptor = st.selectbox("Select a preceptor to spotlight:", preceptor_list)
            
            # 2) If they pick a different name, clear out the old summary
            if st.session_state["_prev_preceptor"] != selected_preceptor:
                st.session_state["_prev_preceptor"] = selected_preceptor
                st.session_state["spotlight_reason"] = ""

            with suppress(NameError):
                # 3) Button to trigger AI
                if st.button("🖋️ Generate Spotlight Summary"):
                    with st.spinner("Generating spotlight summary…"):
                        sel = eligible_df[eligible_df["Evaluator"] == selected_preceptor].iloc[0]
                        st.session_state["spotlight_reason"] = generate_spotlight_summary(
                            sel["strengths_preceptor"],
                            sel["Evaluator"],
                        )
                    st.success("Spotlight Ready!")
                
                # 4) Display only if we have text
                if st.session_state["spotlight_reason"]:
                    # 1) Pull your DF from session_state
                    df_final = st.session_state["df_final_summarized"]
                
                    # 2) Create a boolean mask for the selected preceptor
                    mask = df_final["Evaluator"] == st.session_state["_prev_preceptor"]
    
                    row = df_final.loc[mask].iloc[0]
                
                    # 3) Assign the session_state text into a new column
                    df_final.loc[mask, "spotlight_summary"] = st.session_state["spotlight_reason"]
                
                    # 4) Update your cached DF so later widgets (or downloads) see it too
                    st.session_state["df_final_summarized"] = df_final
                
                    # 5) Extract just that row & display
                    df_spotlight = df_final[mask]
                    st.dataframe(df_spotlight)
                 
                    # --- STEP 3: Upload the Spotlight Record to Firebase ---
                    # Use the evaluator's name as the document ID.
                    record = {
                        "Evaluator": row["Evaluator"],
                        "Evaluator Email": row["Evaluator Email"],
                        "Form Record": str(row["Form Record"]),
                        "spotlight_summary": st.session_state["spotlight_reason"],
                        "Rotation Period": row["Rotation Period"],
                        "num_evaluations": int(row["num_evaluations"]),
                        "strengths_preceptor": row["strengths_preceptor"],
                        "improvement_preceptor": row["improvement_preceptor"],
                    }
    
                ###########################################################################################
                db.collection("spotlight").document(row["Evaluator"]).set(record)
                st.success(f"Spotlight uploaded for {row['Evaluator']}")
                ###########################################################################################
            with suppress(NameError):
                # --- STEP 4: Create a Word Document for the Spotlight Candidate ---
                for idx, row in df_spotlight.iterrows():
                    document = docx.Document()
                    paragraph = document.add_paragraph("PRECEPTOR SPOTLIGHT SUMMARY")
                        
                    # Format the run
                    run = paragraph.runs[0]
                    run.bold = True
                    run.underline = True
                        
                    # Center the paragraph
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph("")
        
                    style = document.styles['Normal']
                    style.font.size = Pt(9)
                    
                    document.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    document.styles['Normal'].paragraph_format.space_before = Pt(0)
                    document.styles['Normal'].paragraph_format.space_after = Pt(0)
        
                    # Create a table with 1 row and 2 columns and a grid style
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    
                    # Optionally disable autofit if available (not always effective)
                    # table.autofit = False
                    
                    # Set each column's cell widths individually:
                    for cell in table.columns[0].cells:
                        set_cell_width(cell, 3)  # 3.45 inches for first column
                    for cell in table.columns[1].cells:
                        set_cell_width(cell, 3.14)  # 2.69 inches for second column
                    
                    # Populate the cells with your content:
                    cell_label = table.rows[0].cells[0]
                    cell_value = table.rows[0].cells[1]
                    
                    # Add bolded "Preceptor:" in the first cell
                    p_label = cell_label.paragraphs[0]
                    run_label = p_label.add_run("Preceptor: ")
                    run_label.bold = True
                    
                    # Add the preceptor name in the second cell (assuming row['Evaluator'] holds the name)
                    p_value = cell_value.paragraphs[0]
                    p_value.add_run(row['Evaluator'])
        
                    shade_cell(cell_label, "ADD8E6")
                    shade_cell(cell_value, "ADD8E6")
        
                    document.add_paragraph("")
        
                    # Create a 6-row, 2-column table
                    details_table = document.add_table(rows=7, cols=2)
                    details_table.style = 'Table Grid'  # Or use your custom border logic
                    
                    # Define column widths if desired
                    col_width_left = Inches(3.0)
                    col_width_right = Inches(3.14)
                    
                    header_row = details_table.rows[0]
                    header_cells = header_row.cells
                    header_cells[0].text = "Evaluation Metric"
                    header_cells[1].text = "Result"
        
                    # Shade the header cells a light gray
                    shade_cell(header_cells[0], "D3D3D3")  # Light gray
                    shade_cell(header_cells[1], "D3D3D3")  # Light gray
                    
                    # Bold the header text
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                                    # Row 1: Rotation Periods
                    details_table.cell(1, 0).text = "Rotation Period (s):"
                    details_table.cell(1, 1).text = str(row['Rotation Period'])
    
                    # Row 2: Number of Evaluations
                    details_table.cell(2, 0).text = "Student-Completed Preceptor Evaluations (n):"
                    details_table.cell(2, 1).text = str(row['num_evaluations'])
                    
                    # Row 3: Number of Student Evaluations Completed by Evaluator
                    details_table.cell(3, 0).text = "Preceptor-Completed Student Evaluations (n):"
                    details_table.cell(3, 1).text = str(row['total_evaluations'])
                    
                    # Row 4: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(4, 0).text = "On-Time Completion Rate (%):"
                    details_table.cell(4, 1).text = f"{row['percentage_on_time']:.1f}%"
    
                    # Row 5: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(5, 0).text = "Number of Students Assigned to Preceptor:"
                    details_table.cell(5, 1).text = str(row['student_assignments'])
    
                    # Row 6: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(6, 0).text = "Number of Times Preceptor Matched to Student:"
                    details_table.cell(6, 1).text = str(row['student_matches'])
    
    
                    # Optionally set column widths
                    for row_idx in range(4):
                        details_table.cell(row_idx, 0).width = col_width_left
                        details_table.cell(row_idx, 1).width = col_width_right
        
                    # Write evaluation question scores.
                    # Assume that the remaining numeric columns (not part of the known text fields) are the evaluation questions.
                    known_cols = {
                        "Evaluator", "Evaluator Email", "Rotation Period", "strengths_preceptor",
                        "improvement_preceptor", "strengths_summary", "improvement_summary",
                        "num_evaluations", "Form Record", "total_evaluations", "percentage_on_time", "student_matches", "student_assignments", "average_prac_score", "average_doc_score", "average_nbme"
                    }
                    document.add_paragraph("")
                    
                    # Define a standard border style
                    border_style = {"sz": "4", "val": "single", "color": "000000"}
                    
                    # Create the table (1 header row, 2 columns)
                    table = document.add_table(rows=1, cols=2)
                    # Do NOT set table.style = 'Table Grid' because we want to manually control borders
                    
                    # Define column widths
                    evaluator_col_width = Inches(6.45)
                    score_col_width = Inches(0.63)
                    
                    # 1) Configure the header row
                    header_row = table.rows[0]
                    header_cells = header_row.cells
                    header_cells[0].text = "Evaluation Question"
                    header_cells[1].text = "Score"
    
                    shade_cell(header_cells[0], "D3D3D3")  # Light gray
                    shade_cell(header_cells[1], "D3D3D3")  # Light gray
                    
                    # Bold the header text
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(9) 
                    
                    # Set the header cell widths
                    header_cells[0].width = evaluator_col_width
                    header_cells[1].width = score_col_width
                    
                    # For the header row:
                    # - keep top/bottom/left/right, but remove the vertical line in the middle
                    set_cell_border(header_cells[0],
                        top=border_style,
                        bottom=border_style,
                        left=border_style
                        # no right border here => removes middle line
                    )
                    set_cell_border(header_cells[1],
                        top=border_style,
                        bottom=border_style,
                        right=border_style
                        # no left border => removes middle line
                    )
                    
                    # 2) Add data rows
                    # Suppose 'question_columns' is a list of the columns that hold numeric data
                    question_columns = [col for col in df_final.columns 
                                        if col not in known_cols 
                                        and pd.api.types.is_numeric_dtype(df_final[col])]
                    
                    for i, col in enumerate(question_columns):
                        row_cells = table.add_row().cells
                        row_cells[0].text = col.rstrip('.')  # question
                        row_cells[1].text = f"{row[col]:.2f}"  # average score
                        
                        # Set widths
                        row_cells[0].width = evaluator_col_width
                        row_cells[1].width = score_col_width
    
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    
                        # For each data row: remove top/bottom to avoid horizontal lines between rows
                        # Keep left and right borders. 
                        # We'll add a bottom border only if it's the last row, to "close" the table.
                        is_last_row = (i == len(question_columns) - 1)
                        
                        set_cell_border(row_cells[0],
                            left=border_style,
                            bottom=border_style if is_last_row else None
                        )
                        set_cell_border(row_cells[1],
                            right=border_style,
                            bottom=border_style if is_last_row else None
                        )
                    
                    document.add_paragraph("")
                    # Create a 3-row, 2-column table
                    details_table = document.add_table(rows=4, cols=2)
                    details_table.style = 'Table Grid'  # Or use your custom border logic
                    
                    # Define column widths if desired
                    col_width_left = Inches(3.0)
                    col_width_right = Inches(3.14)
                    
                    header_row = details_table.rows[0]
                    header_cells = header_row.cells
                    header_cells[0].text = "Assessment Metric (Student Performance)"
                    header_cells[1].text = "Result"
    
                    # Shade the header cells a light gray
                    shade_cell(header_cells[0], "D3D3D3")  # Light gray
                    shade_cell(header_cells[1], "D3D3D3")  # Light gray
                    
                    # Bold the header text
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Row 4: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(1, 0).text = "Practical Examination Score Average(%):"
                    details_table.cell(1, 1).text = f"{row['average_prac_score']:.1f}%"
    
                    # Row 5: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(2, 0).text = "Documentation Score Average (%)"
                    details_table.cell(2, 1).text = f"{row['average_doc_score']:.1f}%"
    
                    # Row 6: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(3, 0).text = "NBME Shelf Score Average (%)"
                    details_table.cell(3, 1).text = f"{row['average_nbme']:.1f}%"
    
    
                    # Optionally set column widths
                    for row_idx in range(4):
                        details_table.cell(row_idx, 0).width = col_width_left
                        details_table.cell(row_idx, 1).width = col_width_right
                    
                    document.add_paragraph("")
                    create_comment_table(document, "Student Documentation Comments", row["documentation_summary"], 6.14)
                    document.add_paragraph("")
                    create_comment_table(document, "Strengths Comments", row["strengths_preceptor"], 6.14)
                    document.add_paragraph("")
                    create_comment_table(document, "Opportunities for Improvement Comments", row["improvement_preceptor"],6.14)
                    document.add_paragraph("")
                    create_comment_table(document, "Spotlight Summary", st.session_state["spotlight_reason"],6.14)
                    
                    # --- STEP 5: Provide a Download Button for the Word Document ---
                    doc_buffer = io.BytesIO()
                    document.save(doc_buffer)
                    doc_buffer.seek(0)
    
                    spotlight_doc_content = doc_buffer.read()

        with suppress(NameError):
            zip_buffer = io.BytesIO()
        
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                zipf.writestr(f"{row['Evaluator'].replace(' ', '_')}_spotlight.docx", spotlight_doc_content)
                # Loop through each evaluator in df_final
                df_final = df_final.loc[df_final['num_evaluations'] >= 1]
                    
                for idx, row in df_final.iterrows():
                    # Create a new Word document for each evaluator
                    document = docx.Document()
                    paragraph = document.add_paragraph("PRECEPTOR PERFORMANCE SUMMARY")
                    
                    # Format the run
                    run = paragraph.runs[0]
                    run.bold = True
                    run.underline = True
                    
                    # Center the paragraph
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
                    document.add_paragraph("")
                    
                    
                    style = document.styles['Normal']
                    style.font.size = Pt(9)
                    
                    document.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    document.styles['Normal'].paragraph_format.space_before = Pt(0)
                    document.styles['Normal'].paragraph_format.space_after = Pt(0)
    
                    # Create a table with 1 row and 2 columns and a grid style
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    
                    # Optionally disable autofit if available (not always effective)
                    # table.autofit = False
                    
                    # Set each column's cell widths individually:
                    for cell in table.columns[0].cells:
                        set_cell_width(cell, 3)  # 3.45 inches for first column
                    for cell in table.columns[1].cells:
                        set_cell_width(cell, 3.14)  # 2.69 inches for second column
                    
                    # Populate the cells with your content:
                    cell_label = table.rows[0].cells[0]
                    cell_value = table.rows[0].cells[1]
                    
                    # Add bolded "Preceptor:" in the first cell
                    p_label = cell_label.paragraphs[0]
                    run_label = p_label.add_run("Preceptor: ")
                    run_label.bold = True
                    
                    # Add the preceptor name in the second cell (assuming row['Evaluator'] holds the name)
                    p_value = cell_value.paragraphs[0]
                    p_value.add_run(row['Evaluator'])
    
                    shade_cell(cell_label, "ADD8E6")
                    shade_cell(cell_value, "ADD8E6")
    
                    document.add_paragraph("")
    
                    # Create a 6-row, 2-column table
                    details_table = document.add_table(rows=7, cols=2)
                    details_table.style = 'Table Grid'  # Or use your custom border logic
                    
                    # Define column widths if desired
                    col_width_left = Inches(3.0)
                    col_width_right = Inches(3.14)
                    
                    header_row = details_table.rows[0]
                    header_cells = header_row.cells
                    header_cells[0].text = "Evaluation Metric"
                    header_cells[1].text = "Result"
    
                    # Shade the header cells a light gray
                    shade_cell(header_cells[0], "D3D3D3")  # Light gray
                    shade_cell(header_cells[1], "D3D3D3")  # Light gray
                    
                    # Bold the header text
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
    
                    # Row 1: Rotation Periods
                    details_table.cell(1, 0).text = "Rotation Period (s):"
                    details_table.cell(1, 1).text = str(row['Rotation Period'])
    
                    # Row 2: Number of Evaluations
                    details_table.cell(2, 0).text = "Student-Completed Preceptor Evaluations (n):"
                    details_table.cell(2, 1).text = str(row['num_evaluations'])
                    
                    # Row 3: Number of Student Evaluations Completed by Evaluator
                    details_table.cell(3, 0).text = "Preceptor-Completed Student Evaluations (n):"
                    details_table.cell(3, 1).text = str(row['total_evaluations'])
                    
                    # Row 4: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(4, 0).text = "On-Time Completion Rate (%):"
                    details_table.cell(4, 1).text = f"{row['percentage_on_time']:.1f}%"
    
                    # Row 5: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(5, 0).text = "Number of Students Assigned to Preceptor:"
                    details_table.cell(5, 1).text = str(row['student_assignments'])
    
                    # Row 6: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(6, 0).text = "Number of Times Preceptor Matched to Student:"
                    details_table.cell(6, 1).text = str(row['student_matches'])
    
    
                    # Optionally set column widths
                    for row_idx in range(4):
                        details_table.cell(row_idx, 0).width = col_width_left
                        details_table.cell(row_idx, 1).width = col_width_right
        
                    # Write evaluation question scores.
                    # Assume that the remaining numeric columns (not part of the known text fields) are the evaluation questions.
                    known_cols = {
                        "Evaluator", "Evaluator Email", "Rotation Period", "strengths_preceptor",
                        "improvement_preceptor", "strengths_summary", "improvement_summary",
                        "num_evaluations", "Form Record", "total_evaluations", "percentage_on_time", "student_matches", "student_assignments", "average_prac_score", "average_doc_score", "average_nbme"
                    }
                    document.add_paragraph("")
                    
                    # Define a standard border style
                    border_style = {"sz": "4", "val": "single", "color": "000000"}
                    
                    # Create the table (1 header row, 2 columns)
                    table = document.add_table(rows=1, cols=2)
                    # Do NOT set table.style = 'Table Grid' because we want to manually control borders
                    
                    # Define column widths
                    evaluator_col_width = Inches(6.45)
                    score_col_width = Inches(0.63)
                    
                    # 1) Configure the header row
                    header_row = table.rows[0]
                    header_cells = header_row.cells
                    header_cells[0].text = "Evaluation Question"
                    header_cells[1].text = "Score"
    
                    shade_cell(header_cells[0], "D3D3D3")  # Light gray
                    shade_cell(header_cells[1], "D3D3D3")  # Light gray
                    
                    # Bold the header text
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(9) 
                    
                    # Set the header cell widths
                    header_cells[0].width = evaluator_col_width
                    header_cells[1].width = score_col_width
                    
                    # For the header row:
                    # - keep top/bottom/left/right, but remove the vertical line in the middle
                    set_cell_border(header_cells[0],
                        top=border_style,
                        bottom=border_style,
                        left=border_style
                        # no right border here => removes middle line
                    )
                    set_cell_border(header_cells[1],
                        top=border_style,
                        bottom=border_style,
                        right=border_style
                        # no left border => removes middle line
                    )
                    
                    # 2) Add data rows
                    # Suppose 'question_columns' is a list of the columns that hold numeric data
                    question_columns = [col for col in df_final.columns 
                                        if col not in known_cols 
                                        and pd.api.types.is_numeric_dtype(df_final[col])]
                    
                    for i, col in enumerate(question_columns):
                        row_cells = table.add_row().cells
                        row_cells[0].text = col.rstrip('.')  # question
                        row_cells[1].text = f"{row[col]:.2f}"  # average score
                        
                        # Set widths
                        row_cells[0].width = evaluator_col_width
                        row_cells[1].width = score_col_width
    
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    
                        # For each data row: remove top/bottom to avoid horizontal lines between rows
                        # Keep left and right borders. 
                        # We'll add a bottom border only if it's the last row, to "close" the table.
                        is_last_row = (i == len(question_columns) - 1)
                        
                        set_cell_border(row_cells[0],
                            left=border_style,
                            bottom=border_style if is_last_row else None
                        )
                        set_cell_border(row_cells[1],
                            right=border_style,
                            bottom=border_style if is_last_row else None
                        )
                    
                    document.add_paragraph("")
                    # Create a 3-row, 2-column table
                    details_table = document.add_table(rows=4, cols=2)
                    details_table.style = 'Table Grid'  # Or use your custom border logic
                    
                    # Define column widths if desired
                    col_width_left = Inches(3.0)
                    col_width_right = Inches(3.14)
                    
                    header_row = details_table.rows[0]
                    header_cells = header_row.cells
                    header_cells[0].text = "Assessment Metric (Student Performance)"
                    header_cells[1].text = "Result"
    
                    # Shade the header cells a light gray
                    shade_cell(header_cells[0], "D3D3D3")  # Light gray
                    shade_cell(header_cells[1], "D3D3D3")  # Light gray
                    
                    # Bold the header text
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Row 4: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(1, 0).text = "Practical Examination Score Average(%):"
                    details_table.cell(1, 1).text = f"{row['average_prac_score']:.1f}%"
    
                    # Row 5: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(2, 0).text = "Documentation Score Average (%)"
                    details_table.cell(2, 1).text = f"{row['average_doc_score']:.1f}%"
    
                    # Row 6: Percentage of Student Evaluations Completed within 14 days
                    details_table.cell(3, 0).text = "NBME Shelf Score Average (%)"
                    details_table.cell(3, 1).text = f"{row['average_nbme']:.1f}%"
    
    
                    # Optionally set column widths
                    for row_idx in range(4):
                        details_table.cell(row_idx, 0).width = col_width_left
                        details_table.cell(row_idx, 1).width = col_width_right
                    
                    document.add_paragraph("")
                    create_comment_table(document, "Student Documentation Comments", row["documentation_summary"], 6.14)
                    document.add_paragraph("")
                    create_comment_table(document, "Strengths Comments", row["strengths_preceptor"], 6.14)
                    document.add_paragraph("")
                    create_comment_table(document, "Opportunities for Improvement Comments", row["improvement_preceptor"],6.14)
                    document.add_paragraph("")
                    create_comment_table(document, "Strengths Summary", row["strengths_summary"],6.14)
                    document.add_paragraph("")
                    create_comment_table(document, "Opportunities for Improvement Summary", row["improvement_summary"],6.14)
    
                    # Save the document to a temporary in-memory buffer
                    doc_buffer = io.BytesIO()
                    document.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    # Create a filename safe for the evaluator (using evaluator's name)
                    safe_name = "".join(c for c in row['Evaluator'] if c.isalnum() or c in (' ', '_')).rstrip().replace(" ", "_")
                    filename = f"{safe_name}.docx"
                    
                    # Write the Word file to the zip archive
                    zipf.writestr(filename, doc_buffer.read())
            
            # Finalize the zip file and get its binary content
            zip_buffer.seek(0)
            zip_data = zip_buffer.getvalue()
    
            st.download_button(label="Download Evaluator Word Files",data=zip_buffer,file_name="evaluators.zip",mime="application/zip",key="download_evaluators_zip")

    except Exception as e:
        st.error(f"Error loading file: {e}")
        


